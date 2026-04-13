"""
Resume text extraction service.
Supports PDF (PyPDF2) and DOCX (python-docx) formats.
"""
import os
import logging

logger = logging.getLogger('resumes')


def parse_resume(file_path: str) -> str:
    """
    Extract text content from a resume file.
    Supports PDF and DOCX formats.
    
    Args:
        file_path: Absolute path to the resume file.
    
    Returns:
        Extracted text as a string.
    
    Raises:
        ValueError: If file format is not supported.
        FileNotFoundError: If file does not exist.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    ext = os.path.splitext(file_path)[1].lower()
    logger.info(f"Parsing resume: {os.path.basename(file_path)} (format: {ext})")

    if ext == '.pdf':
        return _parse_pdf(file_path)
    elif ext in ('.docx', '.doc'):
        return _parse_docx(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}. Please upload PDF or DOCX.")


def _parse_pdf(file_path: str) -> str:
    """Extract text from a PDF file using PyPDF2."""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        text = '\n'.join(text_parts)
        logger.info(f"PDF parsed: {len(text)} characters extracted from {len(reader.pages)} pages")
        return text.strip()
    except ImportError:
        logger.error("PyPDF2 is not installed")
        raise ImportError("PyPDF2 is required for PDF parsing. Install with: pip install PyPDF2")
    except Exception as e:
        logger.error(f"PDF parsing error: {e}")
        raise ValueError(f"Could not parse PDF: {str(e)}")


def _parse_docx(file_path: str) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document
        doc = Document(file_path)
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text.strip())
        text = '\n'.join(text_parts)
        logger.info(f"DOCX parsed: {len(text)} characters extracted")
        return text.strip()
    except ImportError:
        logger.error("python-docx is not installed")
        raise ImportError("python-docx is required for DOCX parsing. Install with: pip install python-docx")
    except Exception as e:
        logger.error(f"DOCX parsing error: {e}")
        raise ValueError(f"Could not parse DOCX: {str(e)}")
